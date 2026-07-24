"""Restricted NousResearch/Hermes Agent runtime for Xiaoce collaboration tasks.

Hermes runs in a dedicated virtual environment and subprocess. The Django
service only exposes a small, auditable toolset: scoped file reads, real
artifact creation, and Hermes' profile-local memory. Shell, network, MCP,
browser, cron, code execution, and sub-agent tools are not enabled.
"""
from __future__ import annotations

import csv
import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import time
import uuid
from pathlib import Path
from typing import Any

from django.conf import settings

from apps.council.llm import (
    credential_status,
    is_model_unavailable_error,
    llm_model_candidates,
    resolve_llm_credentials,
)

from .attachments import MAX_ATTACH_BYTES, store_generated_attachment
from .cancellation import AgentRunCancelled, raise_if_cancelled
from .generated_artifacts import SUPPORTED_EXTENSIONS, sanitize_static_html


logger = logging.getLogger(__name__)

_ARTIFACT_REQUESTS_DIR = ".artifact-requests"
_BINARY_ARTIFACT_EXTENSIONS = {".xlsx", ".docx", ".pdf"}
_ALL_ARTIFACT_EXTENSIONS = SUPPORTED_EXTENSIONS | _BINARY_ARTIFACT_EXTENSIONS
_IMAGE_REQUEST_RE = re.compile(r"(生成|制作|画|绘制).{0,8}(图片|图像|插画|海报)|文生图", re.IGNORECASE)
_SYSTEM_PROMPT = """\
你是良策 AI 协作工作台的小策bot，运行于 NousResearch/Hermes Agent。

工作规则：
1. 默认使用简洁、准确的中文回答；对事实、计算和文件内容负责。
2. 当前工作目录是本会话的隔离空间。input/ 是用户上传或引用的文件，只读；
   artifacts/ 是你的真实文件产物目录。
   用户消息中的【项目知识库与应用证据】来自良策项目知识库或已授权应用，
   与 input/、artifacts/ 是不同来源。只要该区块存在，就必须优先使用；
   不得因为 input/ 或 artifacts/ 为空而声称项目知识库没有资料。
   回答时应明确说明资料来自项目知识库或相应授权来源。
3. 用户要求报告、表格、文档、网页、PDF 或其他文件时，必须真正创建文件，
   不能只在回复中说“已生成”，必须调用 liangce_create_artifact。
4. liangce_create_artifact 的 content：
   - XLSX 使用 JSON：{"sheets":[{"name":"数据","rows":[["列1","列2"],["值1","值2"]]}]}
   - DOCX/PDF 使用完整 Markdown 或纯文本正文。
   - HTML 必须是自包含静态页面，不使用脚本、表单、iframe、远程资源或外链。
5. 可先用 liangce_list_files/liangce_read_file 阅读 input/ 和已有 artifacts/，
   再使用 liangce_create_artifact 生成或覆盖文件。
6. 最终回复说明完成了什么并列出真实文件名；工具失败时如实说明，不要伪造结果。
7. 不执行系统命令、不访问网络、不尝试读取工作区之外的路径。
8. 需要长期记住用户偏好或项目事实时可使用 Hermes memory；不要把密钥写入记忆。
"""


def should_use_hermes(message: str) -> bool:
    """Keep dedicated image generation on the existing image-model path."""
    return not bool(_IMAGE_REQUEST_RE.search(message or ""))


def _safe_filename(value: str, *, default: str = "AI产物.txt") -> str:
    name = Path((value or "").replace("\\", "/")).name
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", name).strip(" .")
    if not name:
        name = default
    if len(name) > 120:
        name = f"{Path(name).stem[:100]}{Path(name).suffix[:12]}"
    return name


def _workspace_for(user_id: int, session_key: str) -> Path:
    base = Path(
        getattr(
            settings,
            "HERMES_WORKSPACE_ROOT",
            settings.BASE_DIR / "hermes_workspaces",
        )
    ).expanduser().resolve()
    digest = hashlib.sha256((session_key or "default").encode("utf-8")).hexdigest()[:20]
    workspace = base / str(user_id) / digest
    (workspace / "input").mkdir(parents=True, exist_ok=True)
    (workspace / "artifacts").mkdir(parents=True, exist_ok=True)
    (workspace / _ARTIFACT_REQUESTS_DIR).mkdir(parents=True, exist_ok=True)
    return workspace


def _path_within(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except (OSError, ValueError):
        return False


class HermesWorkspacePolicy:
    """Path policy shared by tests and the isolated Hermes worker."""

    def __init__(self, workspace: Path) -> None:
        self.workspace = workspace.resolve()
        self.inputs = (workspace / "input").resolve()
        self.artifacts = (workspace / "artifacts").resolve()

    def resolve_read(self, relative_path: str) -> Path:
        raw = Path(str(relative_path or "").replace("\\", "/"))
        if raw.is_absolute() or any(part in {"", ".", ".."} for part in raw.parts):
            raise ValueError("只能使用工作区内的相对路径")
        candidate = (self.workspace / raw).resolve()
        if (
            candidate.name.startswith(".")
            or not _path_within(candidate, self.workspace)
            or not (
                _path_within(candidate, self.inputs)
                or _path_within(candidate, self.artifacts)
            )
        ):
            raise ValueError("只能读取 input/ 或 artifacts/ 内的文件")
        return candidate

    def resolve_artifact(self, filename: str) -> Path:
        safe_name = _safe_filename(filename)
        candidate = (self.artifacts / safe_name).resolve()
        if not _path_within(candidate, self.artifacts):
            raise ValueError("只能向 artifacts/ 写入产物")
        return candidate


def _xlsx_rows(content: str) -> list[tuple[str, list[list[Any]]]]:
    try:
        payload = json.loads(content)
    except json.JSONDecodeError:
        payload = None

    raw_sheets = payload.get("sheets") if isinstance(payload, dict) else None
    sheets: list[tuple[str, list[list[Any]]]] = []
    if isinstance(raw_sheets, list):
        for index, item in enumerate(raw_sheets[:12], 1):
            if not isinstance(item, dict) or not isinstance(item.get("rows"), list):
                continue
            name = str(item.get("name") or f"Sheet{index}")[:31]
            rows = [row[:80] for row in item["rows"][:10_000] if isinstance(row, list)]
            sheets.append((name, rows))
    elif isinstance(payload, list):
        rows = [row[:80] for row in payload[:10_000] if isinstance(row, list)]
        sheets.append(("数据", rows))
    else:
        rows = list(csv.reader(content.splitlines()))
        sheets.append(("数据", rows[:10_000]))
    return sheets or [("数据", [["内容"], [content[:20_000]]])]


def _safe_cell(value: Any) -> Any:
    if isinstance(value, (int, float, bool)) or value is None:
        return value
    text = str(value)[:32_000]
    if text.startswith(("=", "+", "-", "@")):
        text = f"'{text}"
    return text


def _write_xlsx(path: Path, content: str) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    workbook = Workbook()
    default = workbook.active
    workbook.remove(default)
    for sheet_name, rows in _xlsx_rows(content):
        name = sheet_name
        suffix = 2
        while name in workbook.sheetnames:
            name = f"{sheet_name[:27]}-{suffix}"
            suffix += 1
        sheet = workbook.create_sheet(name)
        for row in rows:
            sheet.append([_safe_cell(value) for value in row])
        if sheet.max_row:
            for cell in sheet[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="315EFB")
                cell.alignment = Alignment(vertical="center")
            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions
        for column in sheet.columns:
            width = min(42, max(10, *(len(str(cell.value or "")) + 2 for cell in column[:80])))
            sheet.column_dimensions[column[0].column_letter].width = width
    workbook.save(path)


def _write_docx(path: Path, content: str) -> None:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            document.add_paragraph()
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^[-*]\s+", line):
            document.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
        else:
            document.add_paragraph(line)
    document.save(path)


def _pdf_lines(content: str, font_name: str, font_size: float, width: float) -> list[str]:
    from reportlab.pdfbase.pdfmetrics import stringWidth

    lines: list[str] = []
    for paragraph in content.splitlines() or [""]:
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for char in paragraph:
            candidate = current + char
            if current and stringWidth(candidate, font_name, font_size) > width:
                lines.append(current)
                current = char
            else:
                current = candidate
        lines.append(current)
    return lines


def _write_pdf(path: Path, content: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "Helvetica"
    page_width, page_height = A4
    margin = 54
    font_size = 10.5
    line_height = 17
    document = canvas.Canvas(str(path), pagesize=A4)
    document.setTitle(path.stem)
    document.setAuthor("良策 AI · 小策bot")
    document.setFont(font_name, font_size)
    y = page_height - margin
    for line in _pdf_lines(content, font_name, font_size, page_width - margin * 2):
        if y < margin:
            document.showPage()
            document.setFont(font_name, font_size)
            y = page_height - margin
        document.drawString(margin, y, line)
        y -= line_height
    document.save()


def _write_artifact_file(path: Path, content: str) -> None:
    ext = path.suffix.lower()
    if ext == ".xlsx":
        _write_xlsx(path, content)
    elif ext == ".docx":
        _write_docx(path, content)
    elif ext == ".pdf":
        _write_pdf(path, content)
    else:
        text = sanitize_static_html(content) if ext in {".html", ".htm"} else content
        path.write_text(text, encoding="utf-8")


def _artifact_request_snapshot(workspace: Path) -> set[str]:
    root = workspace / _ARTIFACT_REQUESTS_DIR
    return {path.name for path in root.glob("*.json") if path.is_file()}


def _render_artifact_requests(workspace: Path, before: set[str]) -> None:
    """Render binary artifact requests emitted by the isolated Hermes worker."""
    root = workspace / _ARTIFACT_REQUESTS_DIR
    policy = HermesWorkspacePolicy(workspace)
    for request_path in sorted(root.glob("*.json")):
        if request_path.name in before:
            continue
        try:
            payload = json.loads(request_path.read_text(encoding="utf-8"))
            filename = _safe_filename(str(payload.get("filename") or ""))
            content = str(payload.get("content") or "")
            path = policy.resolve_artifact(filename)
            if path.suffix.lower() not in _BINARY_ARTIFACT_EXTENSIONS:
                raise ValueError("二进制产物请求仅支持 XLSX、DOCX 和 PDF")
            _write_artifact_file(path, content)
            if path.stat().st_size > MAX_ATTACH_BYTES:
                path.unlink(missing_ok=True)
                raise ValueError("产物超过 20MB 限制")
        except Exception:
            logger.exception("Hermes artifact request failed: %s", request_path.name)
        finally:
            request_path.unlink(missing_ok=True)


def _copy_source_files(workspace: Path, source_files: list[dict] | None) -> list[dict]:
    copied: list[dict] = []
    for index, item in enumerate(source_files or []):
        source = Path(str(item.get("path") or "")).expanduser()
        if not source.is_file() or source.stat().st_size > MAX_ATTACH_BYTES:
            continue
        name = _safe_filename(str(item.get("name") or source.name), default=f"input-{index + 1}.txt")
        prefix = re.sub(r"\W+", "", str(item.get("source_id") or ""))[-12:]
        destination = workspace / "input" / (f"{prefix}-{name}" if prefix else name)
        shutil.copyfile(source, destination)
        extracted = str(item.get("text") or "").strip()
        extracted_path = None
        if extracted and source.suffix.lower() not in SUPPORTED_EXTENSIONS:
            extracted_path = destination.with_name(f"{destination.name}.extracted.txt")
            extracted_path.write_text(extracted[:120_000], encoding="utf-8")
        copied.append({
            "name": name,
            "relative_path": str(destination.relative_to(workspace)),
            "extracted_path": (
                str(extracted_path.relative_to(workspace)) if extracted_path else ""
            ),
        })
    return copied


def _artifact_snapshot(root: Path) -> dict[str, tuple[int, int, str]]:
    snapshot: dict[str, tuple[int, int, str]] = {}
    for path in root.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in _ALL_ARTIFACT_EXTENSIONS:
            continue
        stat = path.stat()
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        snapshot[str(path.relative_to(root))] = (stat.st_size, stat.st_mtime_ns, digest)
    return snapshot


def _materialize_changed_artifacts(
    *,
    root: Path,
    before: dict[str, tuple[int, int, str]],
    user_id: int,
) -> list[dict]:
    generated: list[dict] = []
    for relative, current in _artifact_snapshot(root).items():
        if before.get(relative) == current:
            continue
        path = root / relative
        if path.stat().st_size > MAX_ATTACH_BYTES:
            continue
        payload = path.read_bytes()
        if path.suffix.lower() in {".html", ".htm"}:
            safe = sanitize_static_html(payload.decode("utf-8", errors="replace"))
            payload = safe.encode("utf-8")
            path.write_bytes(payload)
        generated.append(
            store_generated_attachment(
                payload=payload,
                filename=path.name,
                user_id=user_id,
                mime=mimetypes.guess_type(path.name)[0] or "application/octet-stream",
            )
        )
    return generated[:8]


def _initial_history(history: list[dict]) -> list[dict]:
    restored: list[dict] = []
    for item in history[-12:]:
        role = "assistant" if item.get("role") == "assistant" else "user"
        text = str(item.get("content") or "").strip()
        if text:
            restored.append({
                "role": role,
                "content": text[:12_000],
            })
    return restored


def _user_prompt(
    *,
    message: str,
    references: list[str],
    inputs: list[dict],
) -> str:
    sections = [message.strip() or "请阅读附件并完成任务。"]
    if inputs:
        listing = []
        for item in inputs:
            detail = f"- {item['name']}: `{item['relative_path']}`"
            if item.get("extracted_path"):
                detail += f"（解析文本：`{item['extracted_path']}`）"
            listing.append(detail)
        sections.append("【本轮可读取文件】\n" + "\n".join(listing))
    if references:
        sections.append(
            "【项目知识库与应用证据（已授权，优先使用）】\n"
            "以下内容来自项目知识库或授权应用，不依赖 input/、artifacts/ 是否有文件：\n"
            + "\n\n".join(str(block)[:18_000] for block in references[:5])
        )
    return "\n\n".join(sections)[:80_000]


def _runtime_python() -> Path:
    configured = str(getattr(settings, "HERMES_RUNTIME_PYTHON", "") or "").strip()
    if configured:
        return Path(os.path.abspath(os.path.expanduser(configured)))
    return Path(os.path.abspath(Path(settings.BASE_DIR) / ".hermes-runtime" / "bin" / "python"))


def _drain_stream_updates(
    stream_path: Path,
    *,
    offset: int,
    pending: str,
    accumulated: str,
    stream_callback=None,
    trace_callback=None,
) -> tuple[int, str, str]:
    """Read complete JSONL deltas without blocking the Hermes worker."""
    if not stream_path.is_file():
        return offset, pending, accumulated
    try:
        with stream_path.open("r", encoding="utf-8") as handle:
            handle.seek(offset)
            chunk = handle.read()
            offset = handle.tell()
    except (OSError, UnicodeError):
        return offset, pending, accumulated
    if not chunk:
        return offset, pending, accumulated
    lines = f"{pending}{chunk}".split("\n")
    pending = lines.pop()
    changed = False
    for raw_line in lines:
        if not raw_line.strip():
            continue
        try:
            event = json.loads(raw_line)
        except (AttributeError, json.JSONDecodeError):
            continue
        if isinstance(event, dict) and event.get("type") == "trace":
            trace = event.get("event")
            if isinstance(trace, dict) and trace_callback:
                try:
                    trace_callback(trace)
                except Exception:
                    logger.exception("Hermes trace callback failed")
            continue
        delta = event.get("delta") if isinstance(event, dict) else None
        if not isinstance(delta, str) or not delta:
            continue
        # Hermes normally emits deltas. Accept cumulative chunks too so a
        # provider adapter cannot duplicate the already-rendered answer.
        accumulated = (
            delta
            if len(delta) >= len(accumulated) and delta.startswith(accumulated)
            else f"{accumulated}{delta}"
        )
        changed = True
    if changed and stream_callback:
        try:
            stream_callback(accumulated)
        except Exception:
            logger.exception("Hermes stream callback failed")
    return offset, pending, accumulated


def _run_worker_process(
    *,
    payload: dict,
    cancel_check,
    timeout: float,
    stream_callback=None,
    trace_callback=None,
) -> dict:
    runtime_python = _runtime_python()
    worker = (Path(settings.BASE_DIR) / "apps" / "core" / "hermes_worker.py").resolve()
    if not runtime_python.is_file():
        raise RuntimeError(
            f"Hermes 隔离运行时不存在：{runtime_python}；"
            "请先执行 scripts/setup-hermes-runtime.sh"
        )
    stream_path: Path | None = None
    worker_payload = dict(payload)
    if stream_callback or trace_callback:
        stream_root = Path(payload["workspace"]) / ".hermes" / "streams"
        stream_root.mkdir(parents=True, exist_ok=True)
        stream_path = stream_root / f"{uuid.uuid4().hex}.jsonl"
        worker_payload["stream_path"] = str(stream_path)
    process = subprocess.Popen(
        [str(runtime_python), str(worker)],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        cwd=str(payload["workspace"]),
        env={
            **os.environ,
            "PYTHONUNBUFFERED": "1",
            "HERMES_HOME": str(Path(payload["workspace"]) / ".hermes"),
            "TERMINAL_CWD": str(payload["workspace"]),
        },
    )
    serialized = json.dumps(worker_payload, ensure_ascii=False)
    started = time.monotonic()
    first_call = True
    stream_offset = 0
    stream_pending = ""
    stream_content = ""
    while True:
        try:
            stdout, stderr = process.communicate(
                input=serialized if first_call else None,
                timeout=0.25,
            )
            break
        except subprocess.TimeoutExpired:
            first_call = False
            if stream_path is not None:
                stream_offset, stream_pending, stream_content = _drain_stream_updates(
                    stream_path,
                    offset=stream_offset,
                    pending=stream_pending,
                    accumulated=stream_content,
                    stream_callback=stream_callback,
                    trace_callback=trace_callback,
                )
            try:
                raise_if_cancelled(cancel_check)
            except AgentRunCancelled:
                process.kill()
                process.communicate()
                if stream_path is not None:
                    stream_path.unlink(missing_ok=True)
                raise
            if time.monotonic() - started >= timeout:
                process.kill()
                process.communicate()
                if stream_path is not None:
                    stream_path.unlink(missing_ok=True)
                raise TimeoutError(f"Hermes Agent 执行超过 {timeout:g} 秒")
    if stream_path is not None:
        _drain_stream_updates(
            stream_path,
            offset=stream_offset,
            pending=stream_pending,
            accumulated=stream_content,
            stream_callback=stream_callback,
            trace_callback=trace_callback,
        )
        stream_path.unlink(missing_ok=True)
    if process.returncode:
        detail = (stderr or stdout or "未知错误").strip()[-2_000:]
        raise RuntimeError(f"Hermes Agent 子进程失败：{detail}")
    try:
        result = json.loads(stdout)
    except json.JSONDecodeError as exc:
        detail = (stdout or stderr or "").strip()[-2_000:]
        raise RuntimeError(f"Hermes Agent 返回格式无效：{detail}") from exc
    if not isinstance(result, dict) or not result.get("ok"):
        raise RuntimeError(str((result or {}).get("error") or "Hermes Agent 未完成任务"))
    return result


def run_hermes_xiaoce(
    *,
    message: str,
    history: list[dict] | None,
    user,
    session_key: str,
    extra_reference_blocks: list[str] | None = None,
    source_files: list[dict] | None = None,
    cancel_check=None,
    progress_callback=None,
    stream_callback=None,
    trace_callback=None,
    agent_name: str = "小策bot",
) -> dict | None:
    """Run Xiaoce in Hermes Agent, returning ``None`` for legacy fallback."""
    if not getattr(settings, "HERMES_AGENT_ENABLED", True) or not should_use_hermes(message):
        return None
    tools_progress_started = False
    try:
        api_key, base_url, model = resolve_llm_credentials(user)
        if not api_key or not base_url or not model:
            raise RuntimeError("Hermes Agent 需要完整的 LLM API Key、Base URL 和模型配置")
        model_status = credential_status(user)
        is_personal_model = model_status.get("source") == "personal"
        if trace_callback:
            provider_label = "DeepSeek" if (
                "deepseek" in model.casefold() or "deepseek" in base_url.casefold()
            ) else "个人模型" if is_personal_model else "平台模型"
            try:
                trace_callback({
                    "id": "model-config",
                    "status": "completed",
                    "label": f"已连接 {provider_label}",
                    "detail": model,
                })
            except Exception:
                logger.exception("Hermes model trace callback failed")
        workspace = _workspace_for(user.id, session_key)
        inputs = _copy_source_files(workspace, source_files)
        artifacts_root = workspace / "artifacts"
        before = _artifact_snapshot(artifacts_root)
        request_before = _artifact_request_snapshot(workspace)
        timeout = max(float(getattr(settings, "HERMES_AGENT_TIMEOUT_SECONDS", 180)), 10)
        max_turns = max(
            1,
            min(int(getattr(settings, "HERMES_AGENT_MAX_TURNS", 10)), 30),
        )
        raise_if_cancelled(cancel_check)
        if progress_callback:
            progress_callback("understanding", "running", {})
            progress_callback("understanding", "completed", {})
            progress_callback("tools", "running", {})
            tools_progress_started = True
        worker_payload = {
            "workspace": str(workspace),
            "api_key": api_key,
            "base_url": base_url,
            "session_key": session_key,
            "user_id": str(user.id),
            "message": _user_prompt(
                message=message,
                references=extra_reference_blocks or [],
                inputs=inputs,
            ),
            "history": _initial_history(history or []),
            "system_prompt": _SYSTEM_PROMPT.replace(
                "的小策bot，",
                f"的{'良策AI' if agent_name == '良策AI' else '小策bot'}，",
                1,
            ),
            "max_turns": max_turns,
        }
        result = None
        last_error: Exception | None = None
        used_model = model
        model_candidates = [model] if is_personal_model else llm_model_candidates(model)
        for candidate in model_candidates:
            try:
                used_model = candidate
                result = _run_worker_process(
                    payload={**worker_payload, "model": candidate},
                    cancel_check=cancel_check,
                    timeout=timeout,
                    stream_callback=stream_callback,
                    trace_callback=trace_callback,
                )
                break
            except AgentRunCancelled:
                raise
            except Exception as exc:
                last_error = exc
                if not is_model_unavailable_error(str(exc)):
                    raise
                logger.warning("Hermes model unavailable, retrying fallback model=%s", candidate)
        if result is None:
            raise last_error or RuntimeError("Hermes Agent 未返回结果")
        _render_artifact_requests(workspace, request_before)
        if progress_callback:
            progress_callback(
                "tools",
                "completed",
                {"tool_count": int(result.get("tool_count") or 0)},
            )
            progress_callback("composing", "running", {})
        generated = _materialize_changed_artifacts(
            root=artifacts_root,
            before=before,
            user_id=user.id,
        )
        if progress_callback:
            progress_callback("composing", "completed", {})
        return {
            "ok": True,
            "reply": str(result.get("reply") or "").strip() or "任务已执行完成。",
            "generated_artifacts": generated,
            "agent_runtime": "hermes-agent",
            "tool_count": int(result.get("tool_count") or 0),
            "hermes_version": str(result.get("hermes_version") or ""),
            "model": used_model,
        }
    except AgentRunCancelled:
        raise
    except Exception:
        logger.exception("Hermes Xiaoce run failed; using stable legacy fallback")
        if stream_callback:
            try:
                stream_callback("")
            except Exception:
                logger.exception("Hermes stream reset callback failed")
        if progress_callback and tools_progress_started:
            progress_callback(
                "tools",
                "completed",
                {"tool_count": 0, "fallback": True},
            )
        return None
